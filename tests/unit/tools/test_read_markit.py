"""Tests for markdown-flavored document/text conversions."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock

import pytest

from kaos.path import KaosPath
from kosong.tooling import ToolOk

from kimi_cli.tools.file.read import Params as ReadFileParams, ReadFile
from kimi_cli.tools.file.read_markit import markdown_to_text


@pytest.fixture
def read_tool(tmp_path: Path) -> ReadFile:
    runtime = MagicMock()
    runtime.builtin_args.KIMI_WORK_DIR = KaosPath(str(tmp_path))
    runtime.additional_dirs = []
    runtime.llm.capabilities = set()
    session = MagicMock()
    session.id = "test"
    session.custom_data = {}
    session.custom_config = {"config_json": {}}
    return ReadFile(runtime, session)


class TestMarkitHelpers:
    def test_markdown_to_text(self) -> None:
        md = "# Title\n\nSome **bold** and `code`.\n\n```python\nx = 1\nx += 1\n```\n\n[A link](http://example.com)\n"
        text = markdown_to_text(md)
        assert "Title" in text
        assert "**" not in text
        assert "`code`" not in text
        assert "[A link]" not in text
        assert "code block:" in text
        assert "example.com" in text

    def test_markdown_to_text_preserves_inline_code_underscores(self) -> None:
        """Inline code must not be rewritten by the italic/emphasis regexes.

        Regression: ``syntax_check.py`` used to become ``syntaxcheck.py`` and
        ``kimi_cli/soul/compaction.py`` became ``kimicli/soul/compaction.py``
        because the `_..._` italic regex paired underscores across code spans.
        """
        md = (
            "- run `uv run tools/syntax_check.py <python_file> ...`\n"
            "- pipeline (`kimi_cli/soul/compaction.py`) + ledger "
            "(`compaction_ledger.py`)\n"
            "- keep `foo_bar` verbatim and _real italic_\n"
        )
        text = markdown_to_text(md)
        assert "syntax_check.py" in text
        assert "<python_file>" in text
        assert "kimi_cli/soul/compaction.py" in text
        assert "compaction_ledger.py" in text
        assert "foo_bar" in text
        assert "syntaxcheck.py" not in text
        assert "kimicli" not in text
        assert "compactionledger" not in text
        assert "real italic" in text


class TestMarkitReadFile:
    async def test_markdown_file_render_markdown(self, read_tool: ReadFile, tmp_path: Path) -> None:
        f = tmp_path / "notes.md"
        f.write_text("# Hello\n\n`code`\n")
        result = await read_tool(ReadFileParams(path=str(f), render_markdown=True))
        assert isinstance(result, ToolOk)
        assert "# Hello" not in result.output
        assert "Hello" in result.output
        assert "`code`" not in result.output

    async def test_html_file_render_markdown(self, read_tool: ReadFile, tmp_path: Path) -> None:
        f = tmp_path / "page.html"
        f.write_text("<h1>Hello</h1><p>World</p>")
        result = await read_tool(ReadFileParams(path=str(f), render_markdown=True))
        assert isinstance(result, ToolOk)
        assert "Hello" in result.output
        assert "World" in result.output

    async def test_render_markdown_false_keeps_raw(self, read_tool: ReadFile, tmp_path: Path) -> None:
        f = tmp_path / "notes.md"
        f.write_text("# Hello\n")
        result = await read_tool(ReadFileParams(path=str(f), render_markdown=False))
        assert isinstance(result, ToolOk)
        assert "# Hello" in result.output

    async def test_docx_markdown_default(self, read_tool: ReadFile, tmp_path: Path) -> None:
        from docx import Document

        f = tmp_path / "report.docx"
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Body text")
        doc.save(str(f))
        result = await read_tool(ReadFileParams(path=str(f)))
        assert isinstance(result, ToolOk)
        assert "## Title" in result.output
        assert "Body text" in result.output
